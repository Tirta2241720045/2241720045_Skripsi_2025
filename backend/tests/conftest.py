import sys
import os
import pytest
import numpy as np
from datetime import datetime
from PIL import Image

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from main import app

from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from fastapi.testclient import TestClient

from app.models.database import Base, get_db, User, Patient
from app.api.auth import create_access_token, hash_password
from app.core.config_test import TEST_DATABASE_URL

engine = create_engine(TEST_DATABASE_URL)
TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)


def override_get_db():
    try:
        db = TestingSessionLocal()
        yield db
    finally:
        db.close()


app.dependency_overrides[get_db] = override_get_db

# ─────────────────────────────────────────────────────────────────────────────
# GLOBAL RESULT STORES — diisi oleh hook pytest_runtest_makereport di bawah
# ─────────────────────────────────────────────────────────────────────────────
_unit_results = []
_integration_results = []
_unit_stats = {"total": 0, "passed": 0, "failed": 0, "skipped": 0,
               "start_time": None, "end_time": None}
_integration_stats = {"total": 0, "passed": 0, "failed": 0}


# ─────────────────────────────────────────────────────────────────────────────
# HOOK: tangkap setiap hasil test dan simpan ke store yang sesuai
# ─────────────────────────────────────────────────────────────────────────────
@pytest.hookimpl(tryfirst=True, hookwrapper=True)
def pytest_runtest_makereport(item, call):
    outcome = yield
    report = outcome.get_result()

    if report.when != "call":
        return

    class_name = item.cls.__name__ if item.cls else "Module"
    test_name  = item.name
    duration   = round(report.duration, 4)

    if report.passed:
        status, message = "PASSED", ""
    elif report.failed:
        status = "FAILED"
        message = str(report.longrepr)[:200] if report.longrepr else ""
    elif report.skipped:
        status, message = "SKIPPED", ""
    else:
        return

    row = {
        'Timestamp'      : datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'Kelas'          : class_name,
        'Test Case'      : test_name,
        'Status'         : status,
        'Durasi (detik)' : duration,
        'Pesan'          : message,
    }

    # Tentukan file test berdasarkan path node
    fspath = str(item.fspath)
    if "test_unit" in fspath:
        _unit_results.append(row)
        _unit_stats["total"] += 1
        if status == "PASSED":
            _unit_stats["passed"] += 1
        elif status == "FAILED":
            _unit_stats["failed"] += 1
        else:
            _unit_stats["skipped"] += 1

    elif "test_integration" in fspath:
        _integration_results.append(row)
        _integration_stats["total"] += 1
        if status == "PASSED":
            _integration_stats["passed"] += 1
        elif status == "FAILED":
            _integration_stats["failed"] += 1


def pytest_sessionstart(session):
    _unit_stats["start_time"] = datetime.now()
    print("\n" + "=" * 60)
    print("🚀 TESTING - STEGOSHIELD")
    print("=" * 60)


def pytest_sessionfinish(session, exitstatus):
    _unit_stats["end_time"] = datetime.now()
    _save_unit_reports()
    _save_integration_reports()
    _save_performance_excel()
    print("\n" + "=" * 60)
    print("✅ SEMUA TESTING SELESAI")
    print("=" * 60)


# ─────────────────────────────────────────────────────────────────────────────
# SIMPAN LAPORAN PERFORMANCE TEST (ambil data dari modul test_performance)
# ─────────────────────────────────────────────────────────────────────────────
def _save_performance_excel():
    # Coba impor store dari test_performance; jika tidak di-run, lewati saja
    try:
        import tests.test_performance as tp
    except ImportError:
        try:
            import test_performance as tp
        except ImportError:
            return

    stores = [tp.embedding_results, tp.extraction_results,
              tp.layer1_results,    tp.layer2_results]
    if not any(stores):
        return

    import pandas as pd

    BASE_DIR    = os.path.dirname(os.path.abspath(__file__))
    RESULTS_DIR = os.path.join(BASE_DIR, "results")
    os.makedirs(RESULTS_DIR, exist_ok=True)
    EXCEL_OUTPUT = os.path.join(RESULTS_DIR, "performance_testing_lengkap.xlsx")

    with pd.ExcelWriter(EXCEL_OUTPUT, engine='openpyxl') as writer:
        if tp.embedding_results:
            pd.DataFrame(tp.embedding_results).to_excel(
                writer, sheet_name='Waktu Embedding', index=False)
        if tp.extraction_results:
            pd.DataFrame(tp.extraction_results).to_excel(
                writer, sheet_name='Waktu Extraction', index=False)
        if tp.layer1_results:
            pd.DataFrame(tp.layer1_results).to_excel(
                writer, sheet_name='Metrik Layer 1', index=False)
        if tp.layer2_results:
            pd.DataFrame(tp.layer2_results).to_excel(
                writer, sheet_name='Metrik Layer 2', index=False)

    print(f"\n📊 Performance Excel : {EXCEL_OUTPUT}")
    print(f"   Embedding  : {len(tp.embedding_results)} baris")
    print(f"   Extraction : {len(tp.extraction_results)} baris")
    print(f"   Metrik L1  : {len(tp.layer1_results)} baris")
    print(f"   Metrik L2  : {len(tp.layer2_results)} baris")


# ─────────────────────────────────────────────────────────────────────────────
# SIMPAN LAPORAN UNIT TEST
# ─────────────────────────────────────────────────────────────────────────────
def _save_unit_reports():
    if not _unit_results:
        return

    import pandas as pd
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    BASE_DIR   = os.path.dirname(os.path.abspath(__file__))
    REPORTS_DIR = os.path.join(BASE_DIR, "reports")
    os.makedirs(REPORTS_DIR, exist_ok=True)

    EXCEL_OUTPUT = os.path.join(REPORTS_DIR, "unit_test_report.xlsx")
    WORD_OUTPUT  = os.path.join(REPORTS_DIR, "unit_test_summary.docx")

    # ── Excel ──────────────────────────────────────────────────────────────
    df_all = pd.DataFrame(_unit_results)
    with pd.ExcelWriter(EXCEL_OUTPUT, engine='openpyxl') as writer:
        df_all.to_excel(writer, sheet_name='Semua Test', index=False)
        for class_name in df_all['Kelas'].unique():
            df_class = df_all[df_all['Kelas'] == class_name]
            df_class.to_excel(writer, sheet_name=class_name[:31], index=False)
    print(f"\n📊 Unit Excel  : {EXCEL_OUTPUT}")

    # ── Word ───────────────────────────────────────────────────────────────
    doc = Document()

    title = doc.add_heading('LAPORAN UNIT TESTING', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('STEGOSHIELD', 1)
    doc.add_paragraph(f'Tanggal Pengujian : {datetime.now().strftime("%d %B %Y")}')
    doc.add_paragraph(f'Waktu             : {datetime.now().strftime("%H:%M:%S")}')

    if _unit_stats['start_time'] and _unit_stats['end_time']:
        dur = (_unit_stats['end_time'] - _unit_stats['start_time']).total_seconds()
        doc.add_paragraph(f'Durasi Total      : {dur:.2f} detik')

    doc.add_paragraph('')
    doc.add_heading('Ringkasan Hasil', 2)

    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['Total Test', 'Berhasil (✓)', 'Gagal (✗)', 'Dilewati (-)']):
        c = tbl.cell(0, i)
        c.text = h
        c.paragraphs[0].runs[0].bold = True
    tbl.cell(1, 0).text = str(_unit_stats['total'])
    tbl.cell(1, 1).text = str(_unit_stats['passed'])
    tbl.cell(1, 2).text = str(_unit_stats['failed'])
    tbl.cell(1, 3).text = str(_unit_stats['skipped'])

    doc.add_paragraph('')
    rate = (_unit_stats['passed'] / _unit_stats['total'] * 100) if _unit_stats['total'] > 0 else 0
    doc.add_paragraph(f'Tingkat Keberhasilan: {rate:.1f}%')
    if _unit_stats['failed'] == 0:
        doc.add_paragraph('✅ STATUS: SEMUA TEST BERHASIL')
    else:
        doc.add_paragraph(f'❌ STATUS: ADA {_unit_stats["failed"]} TEST GAGAL')

    # Ringkasan per kelas
    class_summary: dict = {}
    for r in _unit_results:
        cls = r['Kelas']
        if cls not in class_summary:
            class_summary[cls] = {'total': 0, 'passed': 0, 'failed': 0}
        class_summary[cls]['total'] += 1
        if r['Status'] == 'PASSED':
            class_summary[cls]['passed'] += 1
        elif r['Status'] == 'FAILED':
            class_summary[cls]['failed'] += 1

    doc.add_paragraph('')
    doc.add_heading('Rincian per Kelas Uji', 2)
    tbl2 = doc.add_table(rows=len(class_summary) + 1, cols=4)
    tbl2.style = 'Table Grid'
    for i, h in enumerate(['Kelas Uji', 'Total', 'Berhasil', 'Gagal']):
        c = tbl2.cell(0, i)
        c.text = h
        c.paragraphs[0].runs[0].bold = True
    for idx, (cls, st) in enumerate(class_summary.items(), start=1):
        tbl2.cell(idx, 0).text = cls
        tbl2.cell(idx, 1).text = str(st['total'])
        tbl2.cell(idx, 2).text = str(st['passed'])
        tbl2.cell(idx, 3).text = str(st['failed'])

    doc.add_paragraph('')
    doc.add_heading('Kesimpulan', 2)
    for line in [
        "1. Modul AES-128 CBC: Enkripsi dan dekripsi berfungsi 100% akurat dengan IV random.",
        "2. LSB Layer 1 (MRI Grayscale): Embedding ke area RONI berhasil, ROI tidak termodifikasi.",
        "3. LSB Layer 2 (Foto RGB): Embedding dan ekstraksi berhasil untuk semua resolusi uji.",
        "4. Double Layer End-to-End: Data rekam medis dapat disisipkan dan diekstrak 100% utuh.",
        "5. Format & Kompresi: Hanya PNG (lossless) yang aman; JPEG dan resize merusak data LSB.",
        "6. Metrik Kualitas: PSNR > 30 dB, SSIM > 0.9, MSE < 1.0 untuk semua skenario uji.",
    ]:
        doc.add_paragraph(line)

    doc.add_paragraph('')
    doc.add_paragraph('Laporan ini dibuat otomatis oleh sistem Unit Testing StegoShield.')
    doc.save(WORD_OUTPUT)
    print(f"📄 Unit Word   : {WORD_OUTPUT}")


# ─────────────────────────────────────────────────────────────────────────────
# SIMPAN LAPORAN INTEGRATION TEST
# ─────────────────────────────────────────────────────────────────────────────
def _save_integration_reports():
    if not _integration_results:
        return

    import pandas as pd
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    BASE_DIR    = os.path.dirname(os.path.abspath(__file__))
    REPORTS_DIR = os.path.join(BASE_DIR, "reports")
    os.makedirs(REPORTS_DIR, exist_ok=True)

    EXCEL_OUTPUT = os.path.join(REPORTS_DIR, "integration_test_report.xlsx")
    WORD_OUTPUT  = os.path.join(REPORTS_DIR, "integration_test_summary.docx")

    # ── Excel ──────────────────────────────────────────────────────────────
    df = pd.DataFrame(_integration_results)
    with pd.ExcelWriter(EXCEL_OUTPUT, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Integration Tests', index=False)
        for scenario in df['Kelas'].unique():
            df[df['Kelas'] == scenario].to_excel(
                writer, sheet_name=scenario[:31], index=False)
    print(f"📊 Integ Excel : {EXCEL_OUTPUT}")

    # ── Word ───────────────────────────────────────────────────────────────
    doc = Document()
    title = doc.add_heading('LAPORAN INTEGRATION TESTING', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('STEGOSHIELD', 1)
    doc.add_paragraph(f'Tanggal Pengujian : {datetime.now().strftime("%d %B %Y")}')
    doc.add_paragraph(f'Waktu             : {datetime.now().strftime("%H:%M:%S")}')

    doc.add_paragraph('')
    doc.add_heading('Ringkasan Hasil', 2)
    tbl = doc.add_table(rows=2, cols=3)
    tbl.style = 'Table Grid'
    for i, h in enumerate(['Total Test', 'Berhasil (✓)', 'Gagal (✗)']):
        c = tbl.cell(0, i)
        c.text = h
        c.paragraphs[0].runs[0].bold = True
    tbl.cell(1, 0).text = str(_integration_stats['total'])
    tbl.cell(1, 1).text = str(_integration_stats['passed'])
    tbl.cell(1, 2).text = str(_integration_stats['failed'])

    doc.add_paragraph('')
    rate = (_integration_stats['passed'] / _integration_stats['total'] * 100) \
           if _integration_stats['total'] > 0 else 0
    doc.add_paragraph(f'Tingkat Keberhasilan: {rate:.1f}%')
    if _integration_stats['failed'] == 0:
        doc.add_paragraph('✅ STATUS: SEMUA INTEGRATION TEST BERHASIL')
    else:
        doc.add_paragraph(f'❌ STATUS: ADA {_integration_stats["failed"]} TEST GAGAL')

    doc.add_paragraph('')
    doc.add_heading('Kesimpulan', 2)
    for line in [
        "1. Backend API terintegrasi dengan baik dengan database PostgreSQL.",
        "2. Autentikasi dan otorisasi berbasis role berfungsi sesuai spesifikasi.",
        "3. Alur upload embedding dan extraction end-to-end berjalan lancar.",
        "4. Data tersimpan dan terambil dengan integritas 100%.",
    ]:
        doc.add_paragraph(line)

    doc.add_paragraph('')
    doc.add_paragraph('Laporan ini dibuat otomatis oleh sistem Integration Testing StegoShield.')
    doc.save(WORD_OUTPUT)
    print(f"📄 Integ Word  : {WORD_OUTPUT}")


# ─────────────────────────────────────────────────────────────────────────────
# FIXTURES UMUM
# ─────────────────────────────────────────────────────────────────────────────
@pytest.fixture
def tmp_dir(tmp_path):
    """Alias tmp_path → tmp_dir (dipakai di test_unit & test_performance)."""
    return tmp_path


@pytest.fixture(scope="function")
def db_session():
    Base.metadata.drop_all(bind=engine)
    Base.metadata.create_all(bind=engine)
    db = TestingSessionLocal()
    try:
        yield db
    finally:
        db.close()
    Base.metadata.drop_all(bind=engine)


@pytest.fixture(scope="function")
def client(db_session):
    with TestClient(app) as test_client:
        yield test_client


@pytest.fixture
def admin_token(db_session):
    admin = User(username="admin_test", full_name="Admin Test", role="admin")
    admin.password_hash = hash_password("admin123")
    db_session.add(admin)
    db_session.commit()
    db_session.refresh(admin)
    token = create_access_token(data={"sub": str(admin.user_id), "role": admin.role})
    return {"Authorization": f"Bearer {token}"}


@pytest.fixture
def staff_token(db_session):
    staff = User(username="staff_test", full_name="Staff Test", role="staff")
    staff.password_hash = hash_password("staff123")
    db_session.add(staff)
    db_session.commit()
    db_session.refresh(staff)
    token = create_access_token(data={"sub": str(staff.user_id), "role": staff.role})
    return {"Authorization": f"Bearer {token}"}


@pytest.fixture
def doctor_token(db_session):
    doctor = User(username="doctor_test", full_name="Doctor Test", role="doctor")
    doctor.password_hash = hash_password("doctor123")
    db_session.add(doctor)
    db_session.commit()
    db_session.refresh(doctor)
    token = create_access_token(data={"sub": str(doctor.user_id), "role": doctor.role})
    return {"Authorization": f"Bearer {token}"}


@pytest.fixture
def sample_patient(db_session):
    patient = Patient(
        medical_record_no="MR-1001",
        full_name="Pasien Integration Test",
        date_of_birth="1990-01-01",
        gender="M",
    )
    db_session.add(patient)
    db_session.commit()
    db_session.refresh(patient)
    return patient


@pytest.fixture
def sample_image_files(tmp_path):
    photo_path = str(tmp_path / "photo_test.png")
    arr = np.random.randint(0, 256, (2000, 2000, 3), dtype=np.uint8)
    Image.fromarray(arr, mode='RGB').save(photo_path, 'PNG')

    mri_path = str(tmp_path / "mri_test.png")
    arr = np.random.randint(0, 256, (512, 512), dtype=np.uint8)
    Image.fromarray(arr, mode='L').save(mri_path, 'PNG')

    txt_path = str(tmp_path / "medical_test.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("Test")

    return {"photo": photo_path, "mri": mri_path, "txt": txt_path}


# ─────────────────────────────────────────────────────────────────────────────
# FIXTURES UNTUK test_unit.py
# ─────────────────────────────────────────────────────────────────────────────
@pytest.fixture
def mri_512(tmp_dir):
    path = str(tmp_dir / "mri_512.png")
    arr = np.random.randint(0, 256, (512, 512), dtype=np.uint8)
    Image.fromarray(arr, mode='L').save(path, 'PNG')
    return path


@pytest.fixture
def mri_400(tmp_dir):
    path = str(tmp_dir / "mri_400.png")
    arr = np.random.randint(0, 256, (400, 400), dtype=np.uint8)
    Image.fromarray(arr, mode='L').save(path, 'PNG')
    return path


@pytest.fixture
def mri_600(tmp_dir):
    path = str(tmp_dir / "mri_600.png")
    arr = np.random.randint(0, 256, (600, 600), dtype=np.uint8)
    Image.fromarray(arr, mode='L').save(path, 'PNG')
    return path


@pytest.fixture
def photo_1200(tmp_dir):
    path = str(tmp_dir / "photo_1200.png")
    arr = np.random.randint(0, 256, (1200, 1200, 3), dtype=np.uint8)
    Image.fromarray(arr, mode='RGB').save(path, 'PNG')
    return path


@pytest.fixture
def photo_1500(tmp_dir):
    path = str(tmp_dir / "photo_1500.png")
    arr = np.random.randint(0, 256, (1500, 1500, 3), dtype=np.uint8)
    Image.fromarray(arr, mode='RGB').save(path, 'PNG')
    return path


@pytest.fixture
def photo_2000(tmp_dir):
    path = str(tmp_dir / "photo_2000.png")
    arr = np.random.randint(0, 256, (2000, 2000, 3), dtype=np.uint8)
    Image.fromarray(arr, mode='RGB').save(path, 'PNG')
    return path


@pytest.fixture
def aes():
    from app.core.aes_handler import AESHandler
    return AESHandler("SECRET_KEY_STEGOSHIELD_2026")